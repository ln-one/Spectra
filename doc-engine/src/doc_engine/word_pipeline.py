from __future__ import annotations

import shutil
import subprocess
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from .models import DocumentProfile

PAGEBREAK_MARKER = "[[PAGEBREAK]]"


def set_style_fonts(style, east_asia: str, latin: str, size_pt: float | None = None, bold=None):
    font = style.font
    font.name = latin
    if size_pt is not None:
        font.size = Pt(size_pt)
    if bold is not None:
        font.bold = bold
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east_asia)


def set_run_fonts(run, east_asia: str, latin: str, size_pt: float | None = None, bold=None):
    run.font.name = latin
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east_asia)


def ensure_style(doc: Document, name: str):
    if name in [s.name for s in doc.styles]:
        return doc.styles[name]
    return doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def ensure_caption_style(doc: Document, profile: DocumentProfile):
    style = ensure_style(doc, "Caption")
    set_style_fonts(style, profile.theme.body_cn_font, profile.theme.body_latin_font, 10.5, bold=False)
    pf = style.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.space_before = Pt(3)
    pf.space_after = Pt(3)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE


def apply_document_styles(doc: Document, profile: DocumentProfile):
    theme = profile.theme
    for section in doc.sections:
        section.top_margin = Cm(theme.page_margin_cm)
        section.bottom_margin = Cm(theme.page_margin_cm)
        section.left_margin = Cm(theme.page_margin_cm)
        section.right_margin = Cm(theme.page_margin_cm)
        section.start_type = WD_SECTION_START.NEW_PAGE

    styles = doc.styles
    ensure_caption_style(doc, profile)
    quote = ensure_style(doc, "Quote")
    body_text = ensure_style(doc, "Body Text")
    toc_heading = ensure_style(doc, "TOC Heading")
    toc1 = ensure_style(doc, "toc 1")
    toc2 = ensure_style(doc, "toc 2")
    toc3 = ensure_style(doc, "toc 3")
    list_para = ensure_style(doc, "List Paragraph")

    set_style_fonts(styles["Normal"], theme.body_cn_font, theme.body_latin_font, theme.body_font_size_pt, bold=False)
    pf = styles["Normal"].paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Pt(24)
    pf.line_spacing = Pt(theme.line_spacing_pt)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.space_before = Pt(0)
    pf.space_after = Pt(4)

    set_style_fonts(body_text, theme.body_cn_font, theme.body_latin_font, theme.body_font_size_pt, bold=False)
    body_pf = body_text.paragraph_format
    body_pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body_pf.first_line_indent = Pt(24)
    body_pf.line_spacing = Pt(theme.line_spacing_pt)
    body_pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    body_pf.space_before = Pt(0)
    body_pf.space_after = Pt(4)

    set_style_fonts(styles["Title"], theme.heading_cn_font, theme.heading_latin_font, theme.title_font_size_pt, bold=True)
    title_pf = styles["Title"].paragraph_format
    title_pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_pf.space_before = Pt(120)
    title_pf.space_after = Pt(18)

    set_style_fonts(toc_heading, theme.heading_cn_font, theme.heading_latin_font, 16, bold=True)
    th_pf = toc_heading.paragraph_format
    th_pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    th_pf.first_line_indent = Pt(0)
    th_pf.space_before = Pt(24)
    th_pf.space_after = Pt(12)
    th_pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

    set_style_fonts(styles["Heading 1"], theme.heading_cn_font, theme.heading_latin_font, theme.heading1_size_pt, bold=True)
    set_style_fonts(styles["Heading 2"], theme.heading_cn_font, theme.heading_latin_font, theme.heading2_size_pt, bold=True)
    set_style_fonts(styles["Heading 3"], theme.heading_cn_font, theme.heading_latin_font, theme.heading3_size_pt, bold=True)
    for style_name, before, after in (("Heading 1", 18, 10), ("Heading 2", 12, 6), ("Heading 3", 10, 4)):
        heading_pf = styles[style_name].paragraph_format
        heading_pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        heading_pf.first_line_indent = Pt(0)
        heading_pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        heading_pf.space_before = Pt(before)
        heading_pf.space_after = Pt(after)
        heading_pf.keep_with_next = True

    set_style_fonts(quote, theme.body_cn_font, theme.body_latin_font, 10.5, bold=False)
    quote_pf = quote.paragraph_format
    quote_pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    quote_pf.first_line_indent = Pt(0)
    quote_pf.space_before = Pt(6)
    quote_pf.space_after = Pt(6)
    quote_pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

    for style, size, line_pt in ((toc1, 12, 20), (toc2, 10.5, 18), (toc3, 10.5, 18)):
        set_style_fonts(style, theme.body_cn_font, theme.body_latin_font, size, bold=False)
        toc_pf = style.paragraph_format
        toc_pf.first_line_indent = Pt(0)
        toc_pf.space_before = Pt(0)
        toc_pf.space_after = Pt(2)
        toc_pf.line_spacing = Pt(line_pt)
        toc_pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY

    set_style_fonts(list_para, theme.body_cn_font, theme.body_latin_font, theme.body_font_size_pt, bold=False)
    lp_pf = list_para.paragraph_format
    lp_pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    lp_pf.first_line_indent = Pt(0)
    lp_pf.left_indent = Pt(18)
    lp_pf.space_before = Pt(0)
    lp_pf.space_after = Pt(4)
    lp_pf.line_spacing = Pt(theme.line_spacing_pt)
    lp_pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY


def build_reference_docx(profile: DocumentProfile, output_path: Path) -> None:
    doc = Document()
    apply_document_styles(doc, profile)
    doc.save(output_path)


def run_pandoc(markdown_path: Path, output_docx: Path, reference_docx: Path, include_toc: bool) -> None:
    command = [
        shutil.which("pandoc") or "pandoc",
        str(markdown_path),
        "--from",
        "gfm",
        "--to",
        "docx",
        "--standalone",
        "--reference-doc",
        str(reference_docx),
        "--output",
        str(output_docx),
    ]
    if include_toc:
        command.append("--toc")
    subprocess.run(command, check=True)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_borders = tbl_pr.first_child_found_in("w:tblBorders")
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = tbl_borders.find(qn(f"w:{edge}"))
        if elem is None:
            elem = OxmlElement(f"w:{edge}")
            tbl_borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "8")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), "606060")


def format_tables(doc: Document):
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(table)
        if table.rows:
            set_repeat_table_header(table.rows[0])
            for cell in table.rows[0].cells:
                set_cell_shading(cell, "F2F2F2")
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    para.paragraph_format.first_line_indent = Pt(0)
                    for run in para.runs:
                        run.bold = True
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for para in cell.paragraphs:
                    if not para.text.strip():
                        continue
                    para.paragraph_format.first_line_indent = Pt(0)
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(0)
                    para.paragraph_format.line_spacing = Pt(18)
                    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY


def set_paragraph_border(paragraph, color: str = "BFBFBF", fill: str = "FAFAFA"):
    ppr = paragraph._p.get_or_add_pPr()
    p_bdr = ppr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        ppr.append(p_bdr)
    for edge in ("top", "left", "bottom", "right"):
        elem = p_bdr.find(qn(f"w:{edge}"))
        if elem is None:
            elem = OxmlElement(f"w:{edge}")
            p_bdr.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "8")
        elem.set(qn("w:space"), "8")
        elem.set(qn("w:color"), color)
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)


def add_footer_page_number(section, profile: DocumentProfile):
    section.different_first_page_header_footer = True
    footer = section.footer
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.first_line_indent = Pt(0)
    para.clear()
    run = para.add_run()
    set_run_fonts(run, profile.theme.body_cn_font, profile.theme.body_latin_font, 10.5, False)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def localize_toc_heading(doc: Document):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "Table of Contents":
            paragraph.clear()
            run = paragraph.add_run("目录")
            set_run_fonts(run, "黑体", "Times New Roman", 16, True)
            paragraph.style = doc.styles["TOC Heading"]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def normalize_headings(doc: Document, profile: DocumentProfile):
    for paragraph in doc.paragraphs:
        if paragraph.style and paragraph.style.name in {"Title", "TOC Heading", "Heading 1", "Heading 2", "Heading 3"}:
            for run in paragraph.runs:
                set_run_fonts(
                    run,
                    profile.theme.heading_cn_font,
                    profile.theme.heading_latin_font,
                    None,
                    True,
                )


def normalize_body_runs(doc: Document, profile: DocumentProfile):
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name in {"Heading 1", "Heading 2", "Heading 3", "Title", "TOC Heading"}:
            continue
        for run in paragraph.runs:
            set_run_fonts(run, profile.theme.body_cn_font, profile.theme.body_latin_font)


def add_table_and_figure_captions(doc: Document, profile: DocumentProfile):
    chapter = 0
    table_counter = defaultdict(int)
    previous_was_placeholder = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if paragraph.style and paragraph.style.name == "Heading 1":
            prefix = text.split(".", 1)[0]
            chapter = int(prefix) if prefix.isdigit() else chapter
        if text in {PAGEBREAK_MARKER, "\\newpage"}:
            paragraph.clear()
            paragraph.add_run().add_break(WD_BREAK.PAGE)
            paragraph.paragraph_format.first_line_indent = Pt(0)
            previous_was_placeholder = False
            continue
        if text.startswith("表 "):
            paragraph.style = doc.styles["Caption"]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            previous_was_placeholder = False
            continue
        if text.startswith("图 "):
            paragraph.style = doc.styles["Caption"]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            previous_was_placeholder = False
            continue
        if "【图位占位" in text:
            paragraph.style = doc.styles["Quote"]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(6)
            set_paragraph_border(paragraph)
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(102, 102, 102)
            previous_was_placeholder = True
            continue
        previous_was_placeholder = False

    for table in doc.tables:
        if chapter == 0:
            chapter = 1
        table_counter[chapter] += 1
        caption = doc.add_paragraph(style="Caption")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = caption.add_run(f"表 {chapter}-{table_counter[chapter]}")
        set_run_fonts(run, profile.theme.body_cn_font, profile.theme.body_latin_font, 10.5, False)
        table._tbl.addnext(caption._p)


def postprocess_docx(docx_path: Path, profile: DocumentProfile) -> None:
    doc = Document(docx_path)
    apply_document_styles(doc, profile)
    localize_toc_heading(doc)
    normalize_headings(doc, profile)
    normalize_body_runs(doc, profile)
    format_tables(doc)
    add_table_and_figure_captions(doc, profile)
    if profile.dynamic_fields:
        for section in doc.sections:
            add_footer_page_number(section, profile)
    doc.save(docx_path)


def render_pdf_if_available(docx_path: Path, pdf_path: Path) -> Path | None:
    soffice = shutil.which("soffice")
    if not soffice:
        return None
    subprocess.run(
        [
            soffice,
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(pdf_path.parent),
            str(docx_path),
        ],
        check=True,
    )
    generated = pdf_path.parent / f"{docx_path.stem}.pdf"
    if generated.exists() and generated != pdf_path:
        generated.replace(pdf_path)
    return pdf_path if pdf_path.exists() else None
